import os
import shutil
import re

# Try to import pypdf and python-docx for deep content scanning
try:
    import pypdf
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

def check_pdf_content(filepath, search_terms):
    if not HAS_PYPDF:
        return None
    try:
        with open(filepath, 'rb') as f:
            reader = pypdf.PdfReader(f)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            text_lower = text.lower()
            for term in search_terms:
                if term in text_lower:
                    return term
    except Exception as e:
        print(f"Error reading PDF content of {filepath}: {e}")
    return None

def check_docx_content(filepath, search_terms):
    if not HAS_DOCX:
        return None
    try:
        doc = docx.Document(filepath)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
        text_lower = text.lower()
        for term in search_terms:
            if term in text_lower:
                return term
    except Exception as e:
        print(f"Error reading DOCX content of {filepath}: {e}")
    return None

def main():
    # Target folders
    invoices_dir = os.path.join("Financial", "Invoices")
    receipts_dir = os.path.join("Financial", "Receipts")
    reports_dir = "Reports"

    # Scan current folder
    current_dir = "."
    files = [f for f in os.listdir(current_dir) if os.path.isfile(os.path.join(current_dir, f))]
    
    moved_count = 0

    print("Scanning directory for PDF and DOCX files...")
    if not HAS_PYPDF or not HAS_DOCX:
        print("Warning: Missing deep scanning dependencies. File content search may be restricted.")

    for filename in files:
        filepath = os.path.join(current_dir, filename)
        name, ext = os.path.splitext(filename)
        ext_lower = ext.lower()
        
        if ext_lower not in ['.pdf', '.docx']:
            continue
            
        filename_lower = filename.lower()
        target_dir = None
        reason = ""

        # 1. Check Filename
        if "invoice" in filename_lower:
            target_dir = invoices_dir
            reason = "name contains 'invoice'"
        elif "receipt" in filename_lower:
            target_dir = receipts_dir
            reason = "name contains 'receipt'"
            
        # 2. Check File Content (if not matched by name)
        if not target_dir:
            matched_term = None
            if ext_lower == '.pdf':
                matched_term = check_pdf_content(filepath, ["invoice", "receipt"])
            elif ext_lower == '.docx':
                matched_term = check_docx_content(filepath, ["invoice", "receipt"])
                
            if matched_term == "invoice":
                target_dir = invoices_dir
                reason = "content contains 'invoice'"
            elif matched_term == "receipt":
                target_dir = receipts_dir
                reason = "content contains 'receipt'"

        # 3. Fallback for other DOCX -> Reports
        if not target_dir and ext_lower == '.docx':
            target_dir = reports_dir
            reason = "other DOCX file"

        # Move the file if a target destination was set
        if target_dir:
            os.makedirs(target_dir, exist_ok=True)
            dest = os.path.join(target_dir, filename)
            try:
                shutil.move(filepath, dest)
                print(f"Moved '{filename}' -> '{target_dir}/' ({reason})")
                moved_count += 1
            except Exception as e:
                print(f"Error moving {filename}: {e}")

    print(f"Organization complete. Moved {moved_count} file(s).")

if __name__ == "__main__":
    main()
